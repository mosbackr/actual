"""Investment memo generation pipeline.

Phase 1: Perplexity research (4 parallel calls)
Phase 2: Claude synthesis (1 call producing markdown memo)
Phase 3: Format PDF + DOCX, upload to S3
"""
import asyncio
import io
import logging
import uuid
from datetime import datetime, timezone

import anthropic
import httpx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.colors import HexColor
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from sqlalchemy import select
from sqlalchemy.orm import selectinload

from app.config import settings
from app.db.session import async_session
from app.models.investment_memo import InvestmentMemo, MemoStatus
from app.models.pitch_analysis import AnalysisReport, PitchAnalysis
from app.models.user import User
from app.services import email_service, s3

logger = logging.getLogger(__name__)

# Deep Thesis brand colors (matches analyst_reports.py)
BRAND_ACCENT = "#F28C28"
BRAND_TEXT = "#1A1A1A"
BRAND_TEXT_SECONDARY = "#6B6B6B"

AGENT_LABELS = {
    "problem_solution": "Problem & Solution",
    "market_tam": "Market & TAM",
    "traction": "Traction",
    "technology_ip": "Technology & IP",
    "competition_moat": "Competition & Moat",
    "team": "Team",
    "gtm_business_model": "GTM & Business Model",
    "financials_fundraising": "Financials & Fundraising",
}


# ── Phase 1: Perplexity Research ─────────────────────────────────────

async def _research_perplexity(system_prompt: str, query: str) -> str:
    """Call Perplexity Sonar Pro API. Returns response text or error message."""
    if not settings.perplexity_api_key:
        return "[Perplexity API key not configured — skipping research]"
    try:
        async with httpx.AsyncClient(timeout=90.0) as client:
            resp = await client.post(
                "https://api.perplexity.ai/chat/completions",
                headers={
                    "Authorization": f"Bearer {settings.perplexity_api_key}",
                    "Content-Type": "application/json",
                },
                json={
                    "model": "sonar-pro",
                    "messages": [
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": query},
                    ],
                    "temperature": 0.1,
                    "max_tokens": 4096,
                },
            )
            if resp.status_code == 200:
                data = resp.json()
                return data["choices"][0]["message"]["content"]
            return f"[Perplexity error: {resp.status_code}]"
    except Exception as e:
        logger.warning("Perplexity research failed: %s", e)
        return f"[Research unavailable: {e}]"


async def _run_research(company_name: str, market_context: str, stage: str) -> dict[str, str]:
    """Run 4 parallel Perplexity research queries. Returns dict of research results."""
    system = "Provide factual, detailed research data. Include specific numbers, dates, company names, and sources where available. Be thorough."

    queries = {
        "recent_news": f"Latest news about {company_name} in the last 6 months: funding announcements, product launches, partnerships, press coverage, hiring activity",
        "competitive_landscape": f"Current competitors to {company_name} in {market_context}: their recent funding rounds, market positioning, strengths and weaknesses, market share",
        "market_data": f"Current market size, growth rate (CAGR), trends, and outlook for {market_context} in 2025-2026. Include TAM/SAM estimates from research firms.",
        "comparable_deals": f"Recent VC investment deals and valuations for {stage} startups in {market_context}. Notable exits, M&A activity, and valuation multiples.",
    }

    tasks = {key: _research_perplexity(system, query) for key, query in queries.items()}
    results = await asyncio.gather(*tasks.values(), return_exceptions=True)

    research = {}
    for key, result in zip(tasks.keys(), results):
        if isinstance(result, Exception):
            research[key] = f"[Research failed: {result}]"
        else:
            research[key] = result

    return research


# ── Phase 2: Claude Synthesis ────────────────────────────────────────

MEMO_SYSTEM_PROMPT = """You are a senior venture capital analyst at a top-tier VC firm. You are writing a formal investment memo for your investment committee.

Write a comprehensive, professional investment memo in markdown format. The memo should be thorough, data-driven, and balanced — highlighting both opportunities and risks.

## Required Sections

1. **Executive Summary** — Investment thesis in 3-4 sentences. Include overall score, key strengths, key risks, and your recommendation (Invest / Pass / Watch).

2. **Company Overview** — What the company does, founding stage, product description, and current status.

3. **Market Opportunity** — TAM/SAM/SOM analysis, market growth drivers, timing assessment. Use both the analysis data and fresh market research provided.

4. **Product & Technology** — Solution description, technical differentiation, IP/defensibility, technical risks.

5. **Competitive Landscape** — Key competitors with specific details, competitive positioning, moat analysis. Incorporate fresh competitive intelligence from research.

6. **Team Assessment** — Founder backgrounds, team composition, key gaps, founder-market fit.

7. **Traction & Financials** — Current metrics, growth trajectory, unit economics, financial projections assessment.

8. **Investment Terms** — Recommended raise amount, valuation context based on comparable deals, use of funds assessment.

9. **Risk Factors** — Top 5-7 risks ranked by severity, with mitigation strategies for each.

10. **Recommendation** — Final verdict: **Invest**, **Pass**, or **Watch**. Include conviction level (High/Medium/Low) and specific conditions or milestones that would change your recommendation.

## Formatting Rules
- Use markdown headers (## for sections, ### for subsections)
- Use bullet points for lists
- Use **bold** for emphasis on key metrics and findings
- Include specific numbers and data points wherever possible
- Be concise but thorough — target 2000-3000 words
- Write in third person, professional tone
- Do NOT include a title — the system will add one"""


async def _synthesize_memo(
    company_name: str,
    analysis: dict,
    reports: list[dict],
    research: dict[str, str],
) -> str:
    """Call Claude to synthesize all data into a markdown investment memo."""
    # Build the analysis context
    reports_text = ""
    for r in reports:
        label = AGENT_LABELS.get(r["agent_type"], r["agent_type"])
        reports_text += f"\n### {label} (Score: {r['score']}/100)\n"
        reports_text += f"**Summary:** {r['summary']}\n\n"
        reports_text += f"{r['report']}\n\n"
        if r.get("key_findings"):
            reports_text += "**Key Findings:**\n"
            for f in r["key_findings"]:
                reports_text += f"- {f}\n"
            reports_text += "\n"

    user_message = f"""# Company: {company_name}

## Analysis Metadata
- **Overall Score:** {analysis['overall_score']}/100
- **Fundraising Likelihood:** {analysis.get('fundraising_likelihood', 'N/A')}%
- **Recommended Raise:** {analysis.get('recommended_raise', 'N/A')}
- **Exit Likelihood:** {analysis.get('exit_likelihood', 'N/A')}%
- **Expected Exit Value:** {analysis.get('expected_exit_value', 'N/A')}
- **Expected Exit Timeline:** {analysis.get('expected_exit_timeline', 'N/A')}

## Executive Summary from Analysis
{analysis.get('executive_summary', 'N/A')}

## Detailed Agent Reports
{reports_text}

## Fresh Market Research

### Recent News & Developments
{research.get('recent_news', 'No data available')}

### Competitive Landscape Update
{research.get('competitive_landscape', 'No data available')}

### Market Data
{research.get('market_data', 'No data available')}

### Comparable Deals & Valuations
{research.get('comparable_deals', 'No data available')}

---

Write the investment memo now. Synthesize ALL the above data — both the analysis reports and the fresh research — into a cohesive, professional document."""

    client = anthropic.AsyncAnthropic(api_key=settings.anthropic_api_key)
    response = await client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=8192,
        system=MEMO_SYSTEM_PROMPT,
        messages=[{"role": "user", "content": user_message}],
    )
    return response.content[0].text


# ── Phase 3: PDF + DOCX Formatting ──────────────────────────────────

def _generate_memo_pdf(company_name: str, content: str) -> bytes:
    """Convert markdown memo to branded PDF using reportlab."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter, topMargin=0.75 * inch, bottomMargin=0.75 * inch)
    styles = getSampleStyleSheet()

    styles.add(ParagraphStyle(
        "MemoTitle", parent=styles["Title"], fontSize=28,
        textColor=HexColor(BRAND_ACCENT), spaceAfter=12,
    ))
    styles.add(ParagraphStyle(
        "MemoSubtitle", parent=styles["Title"], fontSize=18,
        textColor=HexColor("#333333"), spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        "MemoDate", parent=styles["Normal"], fontSize=12,
        textColor=HexColor("#808080"), alignment=1, spaceAfter=24,
    ))
    styles.add(ParagraphStyle(
        "MemoH2", parent=styles["Heading2"], fontSize=16,
        textColor=HexColor(BRAND_TEXT), spaceBefore=14, spaceAfter=8,
    ))
    styles.add(ParagraphStyle(
        "MemoH3", parent=styles["Heading3"], fontSize=13,
        textColor=HexColor("#333333"), spaceBefore=10, spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        "MemoBody", parent=styles["BodyText"], fontSize=10,
        leading=14, spaceAfter=8,
    ))
    styles.add(ParagraphStyle(
        "MemoBullet", parent=styles["BodyText"], fontSize=10,
        leading=14, leftIndent=20, bulletIndent=10, spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        "MemoFooter", parent=styles["Normal"], fontSize=8,
        textColor=HexColor("#808080"), alignment=1,
    ))

    story = []

    # Cover page
    story.append(Spacer(1, 2 * inch))
    story.append(Paragraph("Deep Thesis", styles["MemoTitle"]))
    story.append(Paragraph("Investment Memo", styles["MemoSubtitle"]))
    story.append(Spacer(1, 0.5 * inch))
    safe_name = company_name.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    story.append(Paragraph(safe_name, styles["MemoSubtitle"]))
    story.append(Paragraph(datetime.now(timezone.utc).strftime("%B %d, %Y"), styles["MemoDate"]))
    story.append(PageBreak())

    # Parse markdown content into paragraphs
    for block in content.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        safe = block.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

        if block.startswith("## "):
            story.append(Paragraph(safe[3:], styles["MemoH2"]))
        elif block.startswith("### "):
            story.append(Paragraph(safe[4:], styles["MemoH3"]))
        elif block.startswith("- ") or block.startswith("* "):
            for line in block.split("\n"):
                line = line.strip()
                if line.startswith("- ") or line.startswith("* "):
                    safe_line = line[2:].replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    story.append(Paragraph(f"\u2022 {safe_line}", styles["MemoBullet"]))
        else:
            story.append(Paragraph(safe, styles["MemoBody"]))

    # Footer
    story.append(Spacer(1, 24))
    story.append(Paragraph("Generated by Deep Thesis | Confidential", styles["MemoFooter"]))

    doc.build(story)
    buf.seek(0)
    return buf.getvalue()


def _generate_memo_docx(company_name: str, content: str) -> bytes:
    """Convert markdown memo to branded DOCX using python-docx."""
    doc = Document()

    # Cover page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Deep Thesis")
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0xF2, 0x8C, 0x28)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Investment Memo")
    run.font.size = Pt(22)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(company_name)
    run.font.size = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(datetime.now(timezone.utc).strftime("%B %d, %Y"))
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()

    # Parse markdown content
    for block in content.split("\n\n"):
        block = block.strip()
        if not block:
            continue

        if block.startswith("## "):
            doc.add_heading(block[3:], level=2)
        elif block.startswith("### "):
            doc.add_heading(block[4:], level=3)
        elif block.startswith("- ") or block.startswith("* "):
            for line in block.split("\n"):
                line = line.strip()
                if line.startswith("- ") or line.startswith("* "):
                    doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(block)

    # Footer
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run("Generated by Deep Thesis | Confidential")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ── Main Entry Point ─────────────────────────────────────────────────

async def run_memo_generation(memo_id: str) -> None:
    """Main entry point for background memo generation."""
    mid = uuid.UUID(memo_id)

    async with async_session() as db:
        try:
            # Load memo + analysis + reports
            result = await db.execute(
                select(InvestmentMemo).where(InvestmentMemo.id == mid)
            )
            memo = result.scalar_one_or_none()
            if not memo:
                logger.error("Memo %s not found", memo_id)
                return

            result = await db.execute(
                select(PitchAnalysis)
                .where(PitchAnalysis.id == memo.analysis_id)
                .options(selectinload(PitchAnalysis.reports))
            )
            analysis = result.scalar_one()

            # Gather analysis data
            analysis_data = {
                "overall_score": analysis.overall_score,
                "fundraising_likelihood": analysis.fundraising_likelihood,
                "recommended_raise": analysis.recommended_raise,
                "exit_likelihood": analysis.exit_likelihood,
                "expected_exit_value": analysis.expected_exit_value,
                "expected_exit_timeline": analysis.expected_exit_timeline,
                "executive_summary": analysis.executive_summary,
            }

            reports_data = []
            market_context = ""
            for r in analysis.reports:
                agent_type = r.agent_type.value if hasattr(r.agent_type, "value") else r.agent_type
                report_dict = {
                    "agent_type": agent_type,
                    "score": r.score,
                    "summary": r.summary or "",
                    "report": r.report or "",
                    "key_findings": r.key_findings or [],
                }
                reports_data.append(report_dict)
                # Extract market context from the market_tam report
                if agent_type == "market_tam" and r.summary:
                    market_context = r.summary

            company_name = analysis.company_name
            if not market_context:
                market_context = f"{company_name} industry"

            # Determine stage string for research queries
            stage = "seed"
            if analysis.recommended_raise:
                raise_str = analysis.recommended_raise.lower()
                if "series a" in raise_str or "$5" in raise_str or "$10" in raise_str:
                    stage = "series_a"
                elif "series b" in raise_str:
                    stage = "series_b"
                elif "pre" in raise_str:
                    stage = "pre_seed"

            # ── Phase 1: Research ──
            memo.status = "researching"
            await db.commit()

            research = await _run_research(company_name, market_context, stage)

            # ── Phase 2: Synthesis ──
            memo.status = "generating"
            await db.commit()

            content = await _synthesize_memo(company_name, analysis_data, reports_data, research)
            memo.content = content

            # ── Phase 3: Formatting ──
            memo.status = "formatting"
            await db.commit()

            pdf_bytes = _generate_memo_pdf(company_name, content)
            docx_bytes = _generate_memo_docx(company_name, content)

            pdf_key = f"memos/{memo.id}/memo.pdf"
            docx_key = f"memos/{memo.id}/memo.docx"

            s3.upload_file(pdf_bytes, pdf_key)
            s3.upload_file(docx_bytes, docx_key)

            memo.s3_key_pdf = pdf_key
            memo.s3_key_docx = docx_key
            memo.status = "complete"
            memo.completed_at = datetime.now(timezone.utc)
            await db.commit()

            # Send email notification
            user_result = await db.execute(select(User).where(User.id == analysis.user_id))
            user = user_result.scalar_one_or_none()
            if user:
                email_service.send_memo_complete(
                    user_email=user.email,
                    user_name=user.name,
                    analysis_id=str(analysis.id),
                    startup_name=analysis.company_name or "Your startup",
                )

            logger.info("Memo %s generated for %s (PDF: %d bytes, DOCX: %d bytes)",
                        memo_id, company_name, len(pdf_bytes), len(docx_bytes))

        except Exception as e:
            logger.error("Memo generation failed for %s: %s", memo_id, e)
            try:
                memo.status = "failed"
                memo.error = str(e)[:500]
                await db.commit()
            except Exception:
                logger.error("Failed to update memo status for %s", memo_id)
